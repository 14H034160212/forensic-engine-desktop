"""Document parsing: PDF / DOCX / TXT -> clean deck text."""
import io, re, logging
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

def _clean(t):
    t = t.replace("\x00", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()

def parse_pdf(data: bytes) -> str:
    # try pdfplumber (better layout), fall back to pypdf
    try:
        import pdfplumber
        out = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                txt = page.extract_text() or ""
                out.append(f"[Slide/Page {i}]\n{txt}")
        text = "\n\n".join(out)
        if len(text.strip()) > 40:
            return _clean(text)
    except Exception:
        pass
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    out = []
    for i, page in enumerate(reader.pages, 1):
        out.append(f"[Slide/Page {i}]\n{page.extract_text() or ''}")
    return _clean("\n\n".join(out))

def parse_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    # also pull table cells (decks often use tables)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean("\n".join(parts))

def parse(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return parse_pdf(data)
    if name.endswith(".docx"):
        return parse_docx(data)
    if name.endswith((".txt", ".md")):
        return _clean(data.decode("utf-8", "ignore"))
    # sniff: PDF magic
    if data[:4] == b"%PDF":
        return parse_pdf(data)
    if data[:2] == b"PK":   # zip -> probably docx
        try:
            return parse_docx(data)
        except Exception:
            pass
    return _clean(data.decode("utf-8", "ignore"))
